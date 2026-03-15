from docx import Document
from pathlib import Path
from src.ingest.models import NormalizedDocument

def load_docx(file_path):
    path = Path(file_path)
    doc = Document(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    return [
        NormalizedDocument(
            source="docx",
            doc_id=path.stem,
            title=path.name,
            created_at=None,
            content=text,
            metadata={"type": "docx"},
        )
    ]